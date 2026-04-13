import logging
import json
import re
import io
import os
import base64
import uuid
import asyncio
from pathlib import Path as FilePath
from difflib import SequenceMatcher
import traceback

from fastapi import FastAPI, HTTPException, Path, UploadFile, File, Body
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, StreamingResponse, JSONResponse

import config

from models.requests import (
    QuestionRequest,
    RetrievalSettings,
    ConversationCreateRequest,
    ConversationRenameRequest
)

from models.responses import QueryResponse

from llms.engine import get_client
from reflection.engine import ReflectionEngine
from vectordb.engine import VectorDBEngine
from llm_chatbot.simple import SimplePipeline

from dotenv import load_dotenv

from db.db_utils import (
    create_conversation,
    add_message,
    get_messages,
    get_conversations_by_user,
    get_conversation_by_user_and_id,
    delete_conversation_by_user_and_id,
    update_conversation_name,
    sync_user_from_auth_db,
    init_token_usage_table,
    record_token_usage,
    get_usage_stats,
    get_admin_feedback_report,
    init_messages_images_column,
    update_last_bot_message,
    init_message_feedback_table,
    upsert_message_feedback,
    delete_message_feedback,
    get_feedback_stats,
    get_batch_feedback_stats,
    get_message_with_context,
    init_question_feedback_table,
    update_question_feedback,
    compute_question_hash,
)

from metadata_extractor.engine import MetaDataFilterEngine
from hyde.engine import HyDEEngine
from agents.google_tool import GoogleTool
from agents.web_reader_tool import WebReaderTool
from agents.opgg_scraper import scrape_opgg_meta
from agents.comp_evaluator import (
    is_comp_eval_request,
    extract_champions_from_text,
    extract_champions_from_image,
    build_known_champions,
    find_similar_comps,
    format_eval_context,
)
from agents.tft_meta_crawler import (
    is_tft_meta_request,
    detect_content_type,
    crawl_tft_meta,
    format_meta_context,
    format_recipe_table,
    format_recipe_card,
    ITEM_RECIPES,
    _is_recipe_query,
    get_cache as get_meta_cache,
)

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_core.documents import Document
from qdrant_client.models import VectorParams, Distance
from qdrant_client.http import models as qdrant_models


load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="RAG Chatbot API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# =========================
# INIT COMPONENTS
# =========================

vectordb_engine = VectorDBEngine(
    qdrant_host=config.QDRANT_HOST,
    qdrant_port=config.QDRANT_PORT,
    qdrant_collection_name=config.QDRANT_COLLECTION_NAME,
    qdrant_api_key=config.QDRANT_API_KEY
)

llm_client = get_client()

filter_pipeline = MetaDataFilterEngine(llm_client)
hyde_engine = HyDEEngine(llm_client)
reflection_engine = ReflectionEngine(llm_client)

simple_pipeline = SimplePipeline(
    vectordb_engine=vectordb_engine,
    filter_pipeline=filter_pipeline,
    reflection_engine=reflection_engine,
    hyde_engine=hyde_engine,
    default_qa_prompt=config.DEFAULT_QA_PROMPT
)

google_tool = GoogleTool()
web_reader_tool = WebReaderTool(max_length=getattr(config, "WEB_READER_MAX_LENGTH", 3000))

try:
    init_token_usage_table()
    init_messages_images_column()
    init_message_feedback_table()
    init_question_feedback_table()
except Exception:
    pass  # Non-critical: table may already exist or DB may be unavailable

# =========================
# ROUTES
# =========================

@app.get("/")
async def index():
    return FileResponse("index.html")


# =========================
# STREAM GENERATOR
# =========================

def _build_user_content(question: str, images=None):
    """Build OpenAI-compatible user message content.
    Returns str for text-only, list for multimodal (vision)."""
    if not images:
        return question
    content = [{"type": "text", "text": question}]
    for img in images:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:{img.media_type};base64,{img.data}"}
        })
    return content


async def _generate_stream(request: QuestionRequest):

    try:
        reasoning_settings = request.reasoning_settings

        if reasoning_settings is None:
            raise HTTPException(status_code=400, detail="Missing reasoning_settings")

        llm_settings = reasoning_settings.llm

        base_system_prompt = request.system_prompt or config.DEFAULT_SYSTEM_PROMPT
        base_system_prompt = base_system_prompt.replace("{lang}", reasoning_settings.language)

        chat_history = [
            msg.model_dump() for msg in request.chat_history
        ] if request.chat_history else []

        messages = [{"role": "system", "content": base_system_prompt}]
        messages.extend(chat_history)

        original_question = request.question
        # Khi gửi ảnh không kèm text, tự thêm prompt yêu cầu mô tả ảnh
        if request.images and not original_question.strip():
            original_question = "Hãy mô tả chi tiết nội dung của ảnh này."
        user_content = _build_user_content(original_question, request.images)

        if not request.mode:
            raise HTTPException(status_code=400, detail="Missing mode")

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

    # =========================
    # CHITCHAT DETECTION
    # =========================
    def is_chitchat(text: str) -> bool:
        """Detect casual/greeting messages that don't need RAG retrieval."""
        t = text.strip().lower()
        # Remove punctuation for matching
        t_clean = re.sub(r'[^\w\s]', '', t)
        chitchat_patterns = [
            r'^(xin chào|chào|hello|hi|hey|helo|hêy|chao|xinchao)\b',
            r'^(chào bạn|chào bot|xin chào bot|hi bot|hello bot)\b',
            r'^(bạn là ai|bạn tên gì|mày là ai|bạn là gì|giới thiệu bản thân)',
            r'^(bạn có thể làm gì|bạn làm được gì|bạn hỗ trợ gì|bạn giúp được gì)',
            r'^(cảm ơn|cám ơn|thank|thanks|tks|ty)\b',
            r'^(ok|oke|okay|được rồi|hiểu rồi|rõ rồi|ừ|uh|uh huh)\b',
            r'^(tốt|hay|giỏi|tuyệt|awesome|great|nice)\b',
            r'^(bạn khỏe không|khỏe không|how are you)',
            r'^(tạm biệt|bye|goodbye|see you)',
        ]
        for pattern in chitchat_patterns:
            if re.search(pattern, t_clean):
                return True
        # Short messages with no TFT keywords are likely chitchat
        tft_keywords = ['tướng', 'đội hình', 'comp', 'trait', 'item', 'trang bị', 'tier',
                        'meta', 'augment', 'portal', 'patch', 'synergy', 'positioning',
                        'op.gg', 'tft', 'teamfight', 'champion', 'ranked', 'gold']
        words = t_clean.split()
        if len(words) <= 4 and not any(kw in t_clean for kw in tft_keywords):
            return True
        return False

    try:

        # =========================
        # CHITCHAT — trả lời trực tiếp, không qua RAG
        # =========================
        if is_chitchat(original_question) and request.mode == "RAG" and not request.images:
            messages.append({"role": "user", "content": user_content})
            async for chunk in simple_pipeline.stream_completion(
                model_name=llm_settings.model,
                llm_client=llm_client,
                messages=messages
            ):
                yield "data: " + json.dumps(chunk) + "\n\n"
            yield "data: " + json.dumps({"type": "done"}) + "\n\n"
            return

        # =========================
        # COMP EVAL — detect trước tất cả mode
        # =========================
        if is_comp_eval_request(original_question):
            yield "data: " + json.dumps({"type": "status", "message": "Đang phân tích đội hình..."}) + "\n\n"

            try:
                yield "data: " + json.dumps({"type": "info", "message": "Đang lấy meta comps từ op.gg..."}) + "\n\n"
                meta_comps = await scrape_opgg_meta()

                known_champs = build_known_champions(meta_comps)
                user_champions = extract_champions_from_text(original_question, known_champs)

                vision_data = {}

                # Vision fallback: nếu không tìm thấy tướng trong text nhưng có ảnh
                if not user_champions and request.images:
                    yield "data: " + json.dumps({
                        "type": "info",
                        "message": "Đang nhận diện tướng từ ảnh chụp màn hình..."
                    }) + "\n\n"

                    vision_data = await asyncio.to_thread(
                        extract_champions_from_image,
                        llm_client, llm_settings.model,
                        request.images, known_champs
                    )
                    user_champions = vision_data.get("champions", [])

                if not user_champions:
                    if request.images:
                        # Không nhận diện được tướng cụ thể → LLM phân tích ảnh trực tiếp
                        yield "data: " + json.dumps({
                            "type": "info",
                            "message": "Đang phân tích ảnh trực tiếp..."
                        }) + "\n\n"

                        messages[0] = {
                            "role": "system",
                            "content": config.compose_system_prompt(
                                base_system_prompt,
                                config.COMP_EVAL_IMAGE_ONLY_PROMPT_BRANCH
                            )
                        }
                        messages.append({"role": "user", "content": user_content})

                        async for chunk in simple_pipeline.stream_completion(
                            model_name=llm_settings.model,
                            llm_client=llm_client,
                            messages=messages
                        ):
                            yield "data: " + json.dumps(chunk) + "\n\n"
                        yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                        return
                    else:
                        yield "data: " + json.dumps({
                            "type": "token",
                            "content": "⚠️ Không nhận ra tên tướng nào. Hãy **gõ tên tướng cụ thể** hoặc **gửi kèm ảnh chụp toàn màn hình**.\n\nVí dụ: `đánh giá đội hình: Thresh, Braum, Kalista, Annie, Sett, Azir, Ryze, Nidalee, Diana`"
                        }) + "\n\n"
                        yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                        return

                yield "data: " + json.dumps({
                    "type": "info",
                    "message": f"✅ Nhận ra {len(user_champions)} tướng: {', '.join(user_champions)}. Đang so sánh với meta..."
                }) + "\n\n"

                similar = find_similar_comps(user_champions, meta_comps, top_k=3)
                eval_context = format_eval_context(user_champions, similar)

                # Thêm thông tin trang bị từ vision vào context
                if vision_data.get("items_observed"):
                    items_lines = ["\n=== TRANG BỊ HIỆN TẠI (từ ảnh chụp) ==="]
                    for champ, items in vision_data["items_observed"].items():
                        if items:
                            items_lines.append(f"  {champ}: {', '.join(items)}")
                    eval_context += "\n".join(items_lines)

                if vision_data.get("board_level"):
                    eval_context += f"\nLevel hiện tại: {vision_data['board_level']}"
                if vision_data.get("gold"):
                    eval_context += f"\nVàng hiện tại: {vision_data['gold']}"
                if vision_data.get("stage"):
                    eval_context += f"\nGiai đoạn: {vision_data['stage']}"

                # Dùng prompt phù hợp
                if request.images:
                    eval_prompt = config.COMP_EVAL_WITH_IMAGE_PROMPT_BRANCH.format(eval_context=eval_context)
                else:
                    eval_prompt = config.COMP_EVAL_PROMPT_BRANCH.format(eval_context=eval_context)

                messages[0] = {
                    "role": "system",
                    "content": config.compose_system_prompt(base_system_prompt, eval_prompt)
                }
                messages.append({"role": "user", "content": user_content})

                async for chunk in simple_pipeline.stream_completion(
                    model_name=llm_settings.model,
                    llm_client=llm_client,
                    messages=messages
                ):
                    yield "data: " + json.dumps(chunk) + "\n\n"

                yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                return

            except Exception as e:
                messages[0] = {"role": "system", "content": base_system_prompt}
                yield "data: " + json.dumps({
                    "type": "info",
                    "message": f"⚠️ Không thể đánh giá đội hình: {e}. Tiếp tục với mode thường..."
                }) + "\n\n"

        # =========================
        # TFT META — crawl live từ tftacademy + op.gg
        # =========================
        if is_tft_meta_request(original_question) and not request.images:
            yield "data: " + json.dumps({"type": "status", "message": "Đang crawl dữ liệu meta TFT..."}) + "\n\n"

            try:
                content_type = detect_content_type(original_question)
                tftacademy_result, opgg_result = await crawl_tft_meta(content_type)

                tftacademy_data = tftacademy_result.get("data", []) if tftacademy_result.get("success") else []
                opgg_data = opgg_result.get("data", []) if opgg_result.get("success") else []

                # Info messages about crawl results
                info_parts = []
                if tftacademy_data:
                    cached_tag = " (cached)" if tftacademy_result.get("from_cache") else ""
                    info_parts.append(f"tftacademy.com: {len(tftacademy_data)} entries{cached_tag}")
                elif tftacademy_result.get("error"):
                    info_parts.append(f"tftacademy.com: ⚠️ {tftacademy_result['error']}")
                if opgg_data:
                    cached_tag = " (cached)" if opgg_result.get("from_cache") else ""
                    info_parts.append(f"op.gg: {len(opgg_data)} entries{cached_tag}")
                elif opgg_result.get("error") and opgg_result["error"] != "N/A for this content type":
                    info_parts.append(f"op.gg: ⚠️ {opgg_result['error']}")

                if info_parts:
                    yield "data: " + json.dumps({"type": "info", "message": "📊 " + " | ".join(info_parts)}) + "\n\n"

                meta_context = format_meta_context(content_type, tftacademy_data, opgg_data, original_question)

                if meta_context:
                    lang = reasoning_settings.language if reasoning_settings else "Vietnamese"
                    meta_sys_prompt = config.TFT_META_SYS_PROMPT_BRANCH.replace("{lang}", lang)
                    meta_qa_prompt = config.TFT_META_QA_PROMPT.format(
                        meta_context=meta_context,
                        query=original_question
                    )

                    messages[0] = {
                        "role": "system",
                        "content": config.compose_system_prompt(base_system_prompt, meta_sys_prompt)
                    }
                    messages.append({"role": "user", "content": meta_qa_prompt})

                    # Sources
                    sources = []
                    if tftacademy_data:
                        from agents.tft_meta_crawler import TFTACADEMY_URLS
                        sources.append({"title": "TFT Academy", "url": TFTACADEMY_URLS.get(content_type, "")})
                    if opgg_data:
                        sources.append({"title": "OP.GG TFT", "url": "https://op.gg/tft/meta-trends/comps"})
                    if sources:
                        yield "data: " + json.dumps({"type": "sources", "data": sources}) + "\n\n"

                    # Inject visual recipe card with item images if this is a recipe query
                    if _is_recipe_query(original_question):
                        recipe_card = format_recipe_card(original_question, config.BACKEND_BASE_URL)
                        if recipe_card:
                            yield "data: " + json.dumps({"type": "token", "content": recipe_card}) + "\n\n"

                    async for chunk in simple_pipeline.stream_completion(
                        model_name=llm_settings.model,
                        llm_client=llm_client,
                        messages=messages
                    ):
                        yield "data: " + json.dumps(chunk) + "\n\n"

                    yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                    return
                else:
                    yield "data: " + json.dumps({
                        "type": "info",
                        "message": "⚠️ Không crawl được dữ liệu meta. Chuyển sang RAG..."
                    }) + "\n\n"

            except Exception as e:
                messages[0] = {"role": "system", "content": base_system_prompt}
                logger.error(f"TFT Meta crawl error: {e}")
                yield "data: " + json.dumps({
                    "type": "info",
                    "message": f"⚠️ Lỗi crawl meta: {e}. Chuyển sang mode thường..."
                }) + "\n\n"

        # =========================
        # RAG MODE
        # =========================

        if request.mode == "RAG":

            retrieval_settings = request.retrieval_settings or RetrievalSettings(
                retrieval_mode="vector",
                use_MMR=False,
                use_reranking=False,
                use_llm_relevant_scoring=False,
                prioritize_table=False
            )

            # Prepare recipe card for injection if this is a recipe query
            _recipe_card_chunk = None
            if _is_recipe_query(original_question):
                _rc = format_recipe_card(original_question, config.BACKEND_BASE_URL)
                if _rc:
                    _recipe_card_chunk = "data: " + json.dumps({"type": "token", "content": "\n\n" + _rc}) + "\n\n"

            async for chunk in simple_pipeline.stream(
                original_question=original_question,
                llm_client=llm_client,
                messages=messages,
                chat_history=chat_history,
                retrieval_settings=retrieval_settings,
                reasoning_settings=reasoning_settings,
                user_content=user_content
            ):
                # Inject recipe card right before the 'done' event
                if _recipe_card_chunk and '"type": "done"' in chunk:
                    yield _recipe_card_chunk
                    _recipe_card_chunk = None
                yield chunk

        # =========================
        # LLM MODE
        # =========================

        elif request.mode == "LLM":

            yield "data: " + json.dumps({
                "type": "status",
                "message": "Processing with LLM..."
            }) + "\n\n"

            messages.append({
                "role": "user",
                "content": user_content
            })

            async for chunk in simple_pipeline.stream_completion(
                model_name=llm_settings.model,
                llm_client=llm_client,
                messages=messages
            ):

                yield "data: " + json.dumps(chunk) + "\n\n"

            yield "data: " + json.dumps({"type": "done"}) + "\n\n"

        # =========================
        # WEB SEARCH MODE
        # =========================

        elif request.mode == "WEB_SEARCH":

            lang = reasoning_settings.language if reasoning_settings else "Vietnamese"
            web_sys_prompt = config.WEB_SEARCH_SYS_PROMPT_BRANCH.replace("{lang}", lang)
            messages[0] = {
                "role": "system",
                "content": config.compose_system_prompt(base_system_prompt, web_sys_prompt)
            }

            yield "data: " + json.dumps({"type": "status", "message": "Đang xử lý..."}) + "\n\n"

            url_pattern = re.compile(r'https?://[^\s\]\)\"\']+')
            found_urls = url_pattern.findall(original_question)
            query_text = url_pattern.sub("", original_question).strip()
            max_results = getattr(config, "WEB_SEARCH_MAX_RESULTS", 3)

            # =====================================================================
            # FLOW A: Có URL → Scrape → Chunk → Embed vào Qdrant → RAG → LLM
            # =====================================================================
            if found_urls:
                yield "data: " + json.dumps({
                    "type": "info",
                    "message": f"Đang crawl {len(found_urls)} trang web..."
                }) + "\n\n"

                all_docs = []
                source_urls = []

                for url in found_urls:
                    # ── op.gg TFT meta comps: dùng structured scraper ──────────────
                    is_opgg_comps = "op.gg/tft" in url and "comp" in url
                    if is_opgg_comps:
                        yield "data: " + json.dumps({
                            "type": "info",
                            "message": f"Đang crawl op.gg với structured parser..."
                        }) + "\n\n"
                        try:
                            comps = await scrape_opgg_meta()
                            source_urls.append(url)
                            from agents.opgg_scraper import format_comp_as_document
                            for comp in comps:
                                doc_text = format_comp_as_document(comp)
                                all_docs.append(Document(
                                    page_content=doc_text,
                                    metadata={"source": url, "file_name": url, "comp_name": comp["name"]}
                                ))
                            yield "data: " + json.dumps({
                                "type": "info",
                                "message": f"✅ Đã parse {len(comps)} đội hình từ op.gg. Đang phân tích..."
                            }) + "\n\n"
                        except Exception as e:
                            yield "data: " + json.dumps({
                                "type": "info",
                                "message": f"⚠️ Không scrape được op.gg: {e}"
                            }) + "\n\n"
                        continue

                    # ── Các URL thông thường: raw text chunking ───────────────────
                    page_result = await web_reader_tool.read_url(url)
                    if not page_result.get("success") or not page_result.get("content", "").strip():
                        err_detail = page_result.get("error", "unknown error")
                        yield "data: " + json.dumps({
                            "type": "info",
                            "message": f"⚠️ Không đọc được: {url} — {err_detail}"
                        }) + "\n\n"
                        continue

                    raw_content = page_result["content"]
                    source_urls.append(url)

                    yield "data: " + json.dumps({
                        "type": "info",
                        "message": f"✅ Đã crawl {url} ({len(raw_content)} ký tự). Đang phân tích..."
                    }) + "\n\n"

                    # Chunk nội dung
                    splitter = RecursiveCharacterTextSplitter(chunk_size=600, chunk_overlap=80)
                    chunks = splitter.split_text(raw_content)

                    for chunk_text in chunks:
                        all_docs.append(Document(
                            page_content=chunk_text,
                            metadata={"source": url, "file_name": url}
                        ))

                if not all_docs:
                    yield "data: " + json.dumps({
                        "type": "error",
                        "message": "Không crawl được nội dung từ các URL. Trang có thể chặn bot."
                    }) + "\n\n"
                    yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                    return

                yield "data: " + json.dumps({
                    "type": "info",
                    "message": f"Đang embedding {len(all_docs)} đoạn văn bản..."
                }) + "\n\n"

                # Xóa data cũ của các URL này trong Qdrant trước khi insert mới
                try:
                    vectordb_engine.qdrant_client.delete(
                        collection_name=config.QDRANT_COLLECTION_NAME,
                        points_selector=qdrant_models.FilterSelector(
                            filter=qdrant_models.Filter(
                                should=[
                                    qdrant_models.FieldCondition(
                                        key="metadata.source",
                                        match=qdrant_models.MatchValue(value=u)
                                    )
                                    for u in source_urls
                                ]
                            )
                        )
                    )
                except Exception:
                    pass

                # Embed và lưu vào Qdrant
                openai_embeddings = OpenAIEmbeddings(
                    openai_api_key=config.LLM_API_KEY,
                    model="text-embedding-ada-002"
                )
                _ensure_collection(vectordb_engine.qdrant_client, config.QDRANT_COLLECTION_NAME)

                vectors = await asyncio.to_thread(
                    openai_embeddings.embed_documents,
                    [d.page_content for d in all_docs]
                )

                from qdrant_client.models import PointStruct
                points = [
                    PointStruct(
                        id=str(uuid.uuid4()),
                        vector={"dense": vec},
                        payload={
                            "page_content": doc.page_content,
                            "metadata": doc.metadata
                        }
                    )
                    for doc, vec in zip(all_docs, vectors)
                ]
                vectordb_engine.qdrant_client.upsert(
                    collection_name=config.QDRANT_COLLECTION_NAME,
                    points=points
                )

                yield "data: " + json.dumps({
                    "type": "info",
                    "message": f"✅ Đã lưu {len(points)} đoạn vào knowledge base. Đang tìm kiếm câu trả lời..."
                }) + "\n\n"

                # RAG: in-memory cosine similarity trên chính vectors đã embed
                actual_query = query_text if query_text else original_question
                query_vec = await asyncio.to_thread(openai_embeddings.embed_query, actual_query)
                import numpy as np
                q = np.array(query_vec, dtype=np.float32)
                q_norm = q / (np.linalg.norm(q) + 1e-9)
                scored = []
                for doc, vec in zip(all_docs, vectors):
                    v = np.array(vec, dtype=np.float32)
                    v_norm = v / (np.linalg.norm(v) + 1e-9)
                    score = float(np.dot(q_norm, v_norm))
                    scored.append((score, doc.page_content))
                scored.sort(key=lambda x: x[0], reverse=True)
                top_chunks = [text for _, text in scored[:10]]
                context = "\n\n".join(top_chunks)

                web_search_prompt = config.DEFAULT_WEB_SEARCH_QA_PROMPT.format(
                    url=", ".join(source_urls),
                    context=context,
                    query=actual_query
                )
                messages.append({"role": "user", "content": web_search_prompt})

                yield "data: " + json.dumps({
                    "type": "sources",
                    "data": [{"title": u, "url": u} for u in source_urls]
                }) + "\n\n"

                async for chunk in simple_pipeline.stream_completion(
                    model_name=llm_settings.model,
                    llm_client=llm_client,
                    messages=messages
                ):
                    yield "data: " + json.dumps(chunk) + "\n\n"

                yield "data: " + json.dumps({"type": "done"}) + "\n\n"

            # =====================================================================
            # FLOW B: Không có URL → DuckDuckGo search → Scrape → Context → LLM
            # =====================================================================
            else:
                yield "data: " + json.dumps({"type": "info", "message": "Đang tìm kiếm DuckDuckGo..."}) + "\n\n"

                search_result = await google_tool.search(query_text or original_question, max_results=max_results)
                ddg_sources = search_result.get("sources", []) if search_result.get("success") else []

                enriched_sources = []
                if ddg_sources:
                    yield "data: " + json.dumps({
                        "type": "info",
                        "message": f"Tìm thấy {len(ddg_sources)} kết quả. Đang đọc nội dung..."
                    }) + "\n\n"
                    for source in ddg_sources:
                        src_url = source.get("url", "")
                        if src_url:
                            page_result = await web_reader_tool.read_url(src_url)
                            source["full_content"] = page_result.get("content", "") if page_result.get("success") else source.get("content", "")
                        enriched_sources.append(source)

                context_parts = []
                for i, src in enumerate(enriched_sources, 1):
                    content = src.get("full_content") or src.get("content", "")
                    if content.strip():
                        context_parts.append(f"[{i}] {src.get('title', f'Source {i}')}\nURL: {src.get('url', '')}\n{content}\n")

                if not context_parts:
                    yield "data: " + json.dumps({
                        "type": "error",
                        "message": "Không tìm thấy nội dung phù hợp. Hãy thử từ khóa khác hoặc cung cấp URL trực tiếp."
                    }) + "\n\n"
                    yield "data: " + json.dumps({"type": "done"}) + "\n\n"
                    return

                web_search_prompt = config.DEFAULT_WEB_SEARCH_QA_PROMPT.format(
                    url="DuckDuckGo search",
                    context="\n---\n".join(context_parts),
                    query=original_question
                )
                messages.append({"role": "user", "content": web_search_prompt})

                yield "data: " + json.dumps({
                    "type": "sources",
                    "data": [{"title": s.get("title", ""), "url": s.get("url", "")} for s in enriched_sources if s.get("url")]
                }) + "\n\n"

                async for chunk in simple_pipeline.stream_completion(
                    model_name=llm_settings.model,
                    llm_client=llm_client,
                    messages=messages
                ):
                    yield "data: " + json.dumps(chunk) + "\n\n"

                yield "data: " + json.dumps({"type": "done"}) + "\n\n"

        else:

            yield "data: " + json.dumps({
                "type": "error",
                "message": f"Invalid mode {request.mode}"
            }) + "\n\n"

    except Exception as e:

        yield "data: " + json.dumps({
            "type": "error",
            "message": str(e)
        }) + "\n\n"


# =========================
# QUERY API
# =========================

@app.post("/query", response_model=QueryResponse)
async def query(request: QuestionRequest):

    return StreamingResponse(
        _generate_stream(request),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# =========================
# HEALTH
# =========================

@app.get("/health")
async def health_check():
    return {"status": "healthy"}


# =========================
# CONVERSATION APIs
# =========================

@app.post("/conversations")
def api_create_conversation(request: ConversationCreateRequest):

    conversation_id = create_conversation(
        request.user_id,
        request.name
    )

    return {"conversation_id": conversation_id}


@app.put("/conversations/{conversation_id}")
def api_rename_conversation(conversation_id: int, request: ConversationRenameRequest):
    update_conversation_name(request.user_id, conversation_id, request.name.strip())
    return {"success": True}


@app.post("/chat/message")
async def chat_message(request: QuestionRequest):

    conversation_id = getattr(request, "conversation_id", None)
    created_by = getattr(request, "created_by", None)

    if conversation_id and created_by:
        images_data = [{"data": img.data, "media_type": img.media_type} for img in (request.images or [])]
        add_message(
            conversation_id,
            request.question,
            created_by,
            "user",
            images=images_data if images_data else None
        )

    bot_answer = ""
    usage_data = {}

    async def bot_stream():

        nonlocal bot_answer, usage_data

        async for chunk in _generate_stream(request):

            if chunk.startswith("data: "):

                try:

                    data = json.loads(chunk[6:])

                    if data.get("type") == "token":
                        bot_answer += data.get("content", "")
                    elif data.get("type") == "usage":
                        usage_data = data

                except:
                    pass

            yield chunk

        if bot_answer.strip() and conversation_id:

            message_id = add_message(
                conversation_id,
                bot_answer,
                0,
                "assistant"
            )

            # Notify frontend of real DB message_id so feedback buttons can be enabled immediately
            if message_id:
                yield f"data: {json.dumps({'type': 'saved', 'message_id': message_id})}\n\n"

            # Auto-cache successful RAG answers to Qdrant for cross-user knowledge sharing
            if request.mode == "RAG" and len(bot_answer.strip()) > 200:
                _auto_cache_to_qdrant(request.question, bot_answer)

        if usage_data and conversation_id:
            model_name = getattr(getattr(getattr(request, "reasoning_settings", None), "llm", None), "model", config.DEFAULT_MODEL_NAME)
            try:
                record_token_usage(
                    created_by,
                    conversation_id,
                    model_name,
                    usage_data.get("prompt_tokens", 0),
                    usage_data.get("completion_tokens", 0),
                )
            except Exception:
                pass

    return StreamingResponse(
        bot_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        }
    )


@app.get("/history/conversations/{conversation_id}")
def api_get_messages(conversation_id: int):

    messages = get_messages(conversation_id)

    return {"messages": messages}


@app.patch("/messages/{conversation_id}/last-bot")
def api_update_last_bot_message(conversation_id: int, body: dict = Body(...)):
    """Update the last bot message content and sources. Returns the real DB message id."""
    content = body.get("content")
    sources = body.get("sources")
    if not content:
        raise HTTPException(status_code=400, detail="Missing content")
    message_id = update_last_bot_message(conversation_id, content, sources=sources)
    return {"success": True, "message_id": message_id}


@app.get("/history/{user_id}/{conversation_id}")
def get_user_conversation(user_id: int, conversation_id: int):

    conversation = get_conversation_by_user_and_id(
        user_id,
        conversation_id
    )

    return {"conversation": conversation}


@app.delete("/history/{user_id}/{conversation_id}")
def delete_user_conversation(user_id: int, conversation_id: int):

    result = delete_conversation_by_user_and_id(
        user_id,
        conversation_id
    )

    return {"success": result}


@app.get("/history/{user_id}")
def get_conversations(user_id: int):

    conversations = get_conversations_by_user(user_id)

    return {"conversations": conversations}


# =========================
# USER SYNC API
# =========================

@app.post("/sync-user")
def sync_user(user_id: int, username: str, first_name: str = None, last_name: str = None, department_id: int = None, position_id: int = 1):
    """
    Sync user from auth_db to vchatbot database
    Called after successful user registration
    """
    try:
        result = sync_user_from_auth_db(user_id, username, first_name, last_name, department_id, position_id)
        if result:
            return {"success": True, "message": "User synced successfully"}
        else:
            return {"success": True, "message": "User already exists"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# =========================
# UPLOAD DOCUMENT
# =========================

def _ensure_collection(client, collection_name: str):
    """Ensure Qdrant collection exists with dense vector field (1536-dim, COSINE)."""
    try:
        info = client.get_collection(collection_name)
        vectors = info.config.params.vectors
        fields = list(vectors.keys()) if isinstance(vectors, dict) else []
        if "dense" not in fields:
            logger.info(f"Collection '{collection_name}' has wrong schema, recreating...")
            client.delete_collection(collection_name)
            client.create_collection(
                collection_name,
                vectors_config={"dense": VectorParams(size=1536, distance=Distance.COSINE)}
            )
            logger.info(f"Collection '{collection_name}' recreated with dense/1536 schema.")
    except Exception:
        logger.info(f"Collection '{collection_name}' not found, creating...")
        client.create_collection(
            collection_name,
            vectors_config={"dense": VectorParams(size=1536, distance=Distance.COSINE)}
        )
        logger.info(f"Collection '{collection_name}' created with dense/1536 schema.")


@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    filename = file.filename
    if not filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    content = await file.read()
    text = ""

    if filename.lower().endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {e}")

    elif filename.lower().endswith(".txt"):
        text = content.decode("utf-8", errors="replace")

    elif filename.lower().endswith(".docx"):
        try:
            import docx as python_docx
            doc_obj = python_docx.Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc_obj.paragraphs if p.text.strip())
        except ImportError:
            raise HTTPException(status_code=400, detail="python-docx not installed on server")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {e}")

    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type '{filename}'. Supported formats: PDF, TXT, DOCX"
        )

    if not text.strip():
        raise HTTPException(status_code=400, detail="No text could be extracted from file")

    # Split into chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = splitter.split_text(text)
    if not chunks:
        raise HTTPException(status_code=400, detail="No chunks generated from file content")

    # Build document objects
    docs = [
        Document(
            page_content=chunk,
            metadata={"source": filename, "file_name": filename}
        )
        for chunk in chunks
    ]

    # OpenAI embeddings (text-embedding-ada-002, 1536-dim)
    openai_embeddings = OpenAIEmbeddings(
        openai_api_key=config.LLM_API_KEY,
        model="text-embedding-ada-002"
    )

    # Ensure collection has correct schema
    _ensure_collection(vectordb_engine.qdrant_client, config.QDRANT_COLLECTION_NAME)

    # Generate embeddings for all chunks
    logger.info(f"Embedding {len(chunks)} chunks from '{filename}'...")
    vectors = await asyncio.to_thread(
        openai_embeddings.embed_documents,
        [doc.page_content for doc in docs]
    )

    # Build Qdrant points
    from qdrant_client.models import PointStruct
    points = [
        PointStruct(
            id=str(uuid.uuid4()),
            vector={"dense": vec},
            payload={
                "page_content": doc.page_content,
                "metadata": doc.metadata
            }
        )
        for doc, vec in zip(docs, vectors)
    ]

    # Upsert to Qdrant
    vectordb_engine.qdrant_client.upsert(
        collection_name=config.QDRANT_COLLECTION_NAME,
        points=points
    )

    logger.info(f"Uploaded {len(points)} chunks from '{filename}' to Qdrant.")
    return JSONResponse({"status": "success", "chunks_uploaded": len(points), "filename": filename})


# =========================
# INGEST META FROM OP.GG
# =========================

@app.post("/ingest-meta")
async def ingest_meta():
    """
    Scrape op.gg TFT meta comps và ingest vào Qdrant.
    Gọi endpoint này để cập nhật dữ liệu meta mới nhất.
    """
    try:
        comps = await scrape_opgg_meta()
        if not comps:
            return JSONResponse({"status": "error", "message": "Không parse được comp nào từ op.gg"}, status_code=500)

        openai_embeddings = OpenAIEmbeddings(
            openai_api_key=config.LLM_API_KEY,
            model="text-embedding-ada-002"
        )
        _ensure_collection(vectordb_engine.qdrant_client, config.QDRANT_COLLECTION_NAME)

        texts = [c["document_text"] for c in comps]
        logger.info(f"Embedding {len(texts)} comp documents...")
        vectors = await asyncio.to_thread(openai_embeddings.embed_documents, texts)

        from qdrant_client.models import PointStruct
        points = [
            PointStruct(
                id=str(uuid.uuid4()),
                vector={"dense": vec},
                payload={
                    "page_content": text,
                    "metadata": {
                        "source": "op.gg/tft/meta-trends/comps",
                        "file_name": "op.gg TFT Meta",
                        "comp_name": comp["name"],
                        "tier": comp["tier"],
                        "top4_rate": comp["top4_rate"],
                        "avg_place": comp["avg_place"]
                    }
                }
            )
            for text, vec, comp in zip(texts, vectors, comps)
        ]

        vectordb_engine.qdrant_client.upsert(
            collection_name=config.QDRANT_COLLECTION_NAME,
            points=points
        )

        tier_summary = {}
        for c in comps:
            tier_summary[c["tier"]] = tier_summary.get(c["tier"], 0) + 1

        logger.info(f"Ingested {len(points)} comp documents. Tiers: {tier_summary}")
        return JSONResponse({
            "status": "success",
            "comps_ingested": len(points),
            "tiers": tier_summary,
            "preview": [{"name": c["name"], "tier": c["tier"], "top4_rate": c["top4_rate"]} for c in comps[:5]]
        })

    except Exception as e:
        logger.error(f"ingest-meta error: {e}")
        return JSONResponse({"status": "error", "message": str(e)}, status_code=500)


@app.get("/upload/stats")
def upload_stats():
    """Return collection statistics."""
    try:
        info = vectordb_engine.qdrant_client.get_collection(config.QDRANT_COLLECTION_NAME)
        return {
            "collection": config.QDRANT_COLLECTION_NAME,
            "points_count": info.points_count,
            "status": str(info.status)
        }
    except Exception as e:
        return {"collection": config.QDRANT_COLLECTION_NAME, "points_count": 0, "error": str(e)}


# =========================
# CHAMPION IMAGES
# =========================

CHAMPION_IMAGES_DIR = FilePath("champion_images")
CHAMPION_IMAGES_DIR.mkdir(exist_ok=True)

@app.post("/save-image")
async def save_image(name: str, body: dict = Body(...)):
    """Lưu ảnh base64 vào thư mục champion_images với tên cho trước."""
    try:
        img_data = base64.b64decode(body["base64"])
        media_type = body.get("media_type", "image/png")
        ext = media_type.split("/")[-1].split(";")[0]  # png, jpeg, webp, gif
        if ext == "jpeg":
            ext = "jpg"
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', name.strip().lower())
        file_path = CHAMPION_IMAGES_DIR / f"{safe_name}.{ext}"
        with open(file_path, "wb") as f:
            f.write(img_data)
        return {"success": True, "message": f"Đã lưu ảnh '{name}'", "name": safe_name}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@app.get("/image/{name}")
async def get_image(name: str):
    """Trả về file ảnh đã lưu theo tên."""
    safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', name.strip().lower())
    for ext in ["png", "jpg", "jpeg", "webp", "gif"]:
        file_path = CHAMPION_IMAGES_DIR / f"{safe_name}.{ext}"
        if file_path.exists():
            return FileResponse(str(file_path), media_type=f"image/{ext}")
    raise HTTPException(status_code=404, detail=f"Không tìm thấy ảnh '{name}'")


@app.get("/saved-images")
async def get_saved_images():
    """Trả về danh sách tên tướng đã có ảnh lưu trong hệ thống."""
    if not CHAMPION_IMAGES_DIR.exists():
        return {"names": []}
    names = [
        f.stem for f in CHAMPION_IMAGES_DIR.iterdir()
        if f.is_file() and f.suffix.lower() in ['.png', '.jpg', '.jpeg', '.webp', '.gif']
    ]
    return {"names": names}


def _feedback_qdrant_id(message_id: int) -> str:
    """Generate a deterministic UUID for a feedback-boosted Qdrant point (legacy, per-message)."""
    import hashlib
    return str(uuid.UUID(hashlib.md5(f"feedback_{message_id}".encode()).hexdigest()))


def _feedback_qdrant_id_by_question(question: str) -> str:
    """Generate a deterministic UUID based on question text so all feedback converges."""
    import hashlib
    normalized = question.strip().lower()
    return str(uuid.UUID(hashlib.md5(f"feedback_q_{normalized}".encode()).hexdigest()))


def _auto_cache_qdrant_id(question: str) -> str:
    """Generate a deterministic UUID based on the question text (deduplication)."""
    import hashlib
    normalized = question.strip().lower()
    return str(uuid.UUID(hashlib.md5(f"auto_cache_{normalized}".encode()).hexdigest()))


def _auto_cache_to_qdrant(question: str, answer: str):
    """Auto-cache a successful RAG Q&A pair to Qdrant for cross-user knowledge sharing.
    Runs in a background thread to avoid blocking the response."""
    import concurrent.futures as _cf
    import threading

    def _do_cache():
        try:
            q_text = question.strip()
            a_text = answer.strip()
            if not q_text or not a_text:
                return

            doc_text = f"Câu hỏi: {q_text}\n\nCâu trả lời: {a_text}"
            qdrant_id = _auto_cache_qdrant_id(q_text)

            openai_embeddings = OpenAIEmbeddings(
                openai_api_key=config.LLM_API_KEY,
                model="text-embedding-ada-002"
            )
            vec = openai_embeddings.embed_query(q_text)
            _ensure_collection(vectordb_engine.qdrant_client, config.QDRANT_COLLECTION_NAME)

            from qdrant_client.models import PointStruct
            vectordb_engine.qdrant_client.upsert(
                collection_name=config.QDRANT_COLLECTION_NAME,
                points=[PointStruct(
                    id=qdrant_id,
                    vector={"dense": vec},
                    payload={
                        "page_content": doc_text,
                        "metadata": {
                            "source": "auto_cached_answer",
                            "file_name": "Câu trả lời đã được xác nhận",
                            "is_auto_cached": True,
                        }
                    }
                )]
            )
            logger.info(f"Auto-cached Q&A to Qdrant: '{q_text[:80]}...'")
        except Exception as e:
            logger.warning(f"Auto-cache to Qdrant failed: {e}")

    threading.Thread(target=_do_cache, daemon=True).start()


@app.post("/messages/{message_id}/feedback")
def api_submit_feedback(message_id: int, body: dict = Body(...)):
    """Submit thumbs-up or thumbs-down for a bot message, or 'none' to remove vote.
    Feedback is accumulated at the question level so it persists across conversations."""
    user_id = body.get("user_id")
    feedback = body.get("feedback")
    if feedback not in ("up", "down", "none"):
        raise HTTPException(status_code=400, detail="feedback must be 'up', 'down', or 'none'")

    # Get old per-message vote BEFORE changing (for delta calculation)
    old_vote = None
    try:
        old_stats = get_feedback_stats(message_id, user_id)
        old_vote = old_stats.get('user_vote')
    except Exception:
        pass

    # Handle vote removal
    if feedback == "none":
        delete_message_feedback(message_id, user_id)
    else:
        # Upsert per-message feedback (audit trail)
        upsert_message_feedback(message_id, user_id, feedback)

    # Update question-level cumulative feedback
    q_text = None
    msg = None
    question = None
    try:
        msg, question = get_message_with_context(message_id)
        if msg and msg.get('role') == 'assistant' and question:
            q_text = question.get('content', '').strip()
            if q_text:
                delta_up = (1 if feedback == 'up' else 0) - (1 if old_vote == 'up' else 0)
                delta_down = (1 if feedback == 'down' else 0) - (1 if old_vote == 'down' else 0)
                update_question_feedback(q_text, delta_up, delta_down)
    except Exception as e:
        logger.warning(f"Question feedback update error: {e}")

    # Get updated stats (now returns cumulative question-level counts)
    stats = get_feedback_stats(message_id, user_id)
    net_score = stats['up'] - stats['down']

    # Sync to Qdrant: use question-based ID so all feedback converges on one point
    try:
        if msg and msg.get('role') == 'assistant' and question and q_text:
            qdrant_id = _feedback_qdrant_id_by_question(q_text)
            if net_score > 0:
                a_text = msg.get('content', '').strip()
                doc_text = f"Câu hỏi: {q_text}\n\nCâu trả lời: {a_text}"
                openai_embeddings = OpenAIEmbeddings(
                    openai_api_key=config.LLM_API_KEY,
                    model="text-embedding-ada-002"
                )
                import concurrent.futures as _cf
                with _cf.ThreadPoolExecutor() as _pool:
                    vec = _pool.submit(openai_embeddings.embed_query, q_text).result()
                _ensure_collection(vectordb_engine.qdrant_client, config.QDRANT_COLLECTION_NAME)
                from qdrant_client.models import PointStruct
                vectordb_engine.qdrant_client.upsert(
                    collection_name=config.QDRANT_COLLECTION_NAME,
                    points=[PointStruct(
                        id=qdrant_id,
                        vector={"dense": vec},
                        payload={
                            "page_content": doc_text,
                            "metadata": {
                                "source": "feedback_community",
                                "file_name": "⭐ Câu trả lời được cộng đồng đánh giá tốt",
                                "is_feedback_boosted": True,
                                "feedback_score": net_score,
                            }
                        }
                    )]
                )
                logger.info(f"Upserted feedback-boosted doc for question (net={net_score})")
            else:
                # Remove from Qdrant if score drops to 0 or below
                try:
                    from qdrant_client.http.models import PointIdsList
                    vectordb_engine.qdrant_client.delete(
                        collection_name=config.QDRANT_COLLECTION_NAME,
                        points_selector=PointIdsList(points=[qdrant_id])
                    )
                except Exception:
                    pass
                # Also clean up legacy message-based points
                try:
                    vectordb_engine.qdrant_client.delete(
                        collection_name=config.QDRANT_COLLECTION_NAME,
                        points_selector=qdrant_models.FilterSelector(
                            filter=qdrant_models.Filter(
                                must=[qdrant_models.FieldCondition(
                                    key="metadata.feedback_message_id",
                                    match=qdrant_models.MatchValue(value=message_id)
                                )]
                            )
                        )
                    )
                except Exception:
                    pass
    except Exception as e:
        logger.warning(f"Qdrant feedback sync error: {e}")

    return stats


@app.get("/messages/{message_id}/feedback")
def api_get_feedback(message_id: int, user_id: int = None):
    """Get thumbs-up/down counts and user's own vote for a message."""
    return get_feedback_stats(message_id, user_id)


@app.post("/messages/feedback/batch")
def api_batch_feedback(body: dict = Body(...)):
    """Get feedback stats for multiple messages at once."""
    message_ids = body.get("message_ids", [])
    user_id = body.get("user_id")
    if not message_ids:
        return {}
    return get_batch_feedback_stats(message_ids, user_id)


@app.get("/report/usage")
def api_usage_stats(user_id: int = None, days: int = 30):
    try:
        stats = get_usage_stats(user_id=user_id, days=days)
        return stats
    except Exception as e:
        return {"summary": {}, "daily": [], "error": str(e)}


@app.get("/report/admin/feedback")
def api_admin_feedback_report(days: int = 30):
    """Admin endpoint: aggregated feedback stats (up/down) with daily breakdown and top messages."""
    try:
        return get_admin_feedback_report(days=days)
    except Exception as e:
        return {"summary": {}, "daily": [], "top_liked": [], "top_disliked": [], "error": str(e)}


# =========================
# META CACHE MANAGEMENT
# =========================

@app.get("/meta-cache/stats")
def meta_cache_stats():
    """Return TFT meta cache statistics."""
    return get_meta_cache().stats()


@app.post("/meta-cache/clear")
def meta_cache_clear():
    """Clear all TFT meta cache entries."""
    get_meta_cache().clear()
    return {"success": True, "message": "Meta cache cleared"}


# =========================
# QDRANT DATA CLEANUP
# =========================

QDRANT_ITEM_REPLACEMENTS = [
    ("Dao Statikk Huyền Thoại", "Thú Tượng Thạch Giáp"),
    ("Artifact Statikk Shiv", "Gargoyle Stoneplate"),
    ("Dao Statikk", "Thú Tượng Thạch Giáp"),
    ("Statikk Shiv", "Thú Tượng Thạch Giáp"),
]


@app.post("/qdrant/cleanup-removed-items")
def cleanup_removed_items():
    """Find and replace removed items (e.g. Statikk Shiv) in Qdrant entries."""
    try:
        updated = []
        offset = None
        while True:
            results = vectordb_engine.qdrant_client.scroll(
                collection_name=config.QDRANT_COLLECTION_NAME,
                limit=100,
                with_payload=True,
                with_vectors=False,
                offset=offset,
            )
            points, next_offset = results
            for p in points:
                content = p.payload.get("page_content", "") or ""
                if "statikk" not in content.lower():
                    continue
                original = content
                for old, new in QDRANT_ITEM_REPLACEMENTS:
                    content = re.sub(re.escape(old), new, content, flags=re.IGNORECASE)
                if content != original:
                    vectordb_engine.qdrant_client.set_payload(
                        collection_name=config.QDRANT_COLLECTION_NAME,
                        payload={"page_content": content},
                        points=[p.id],
                    )
                    src = p.payload.get("metadata", {}).get("source", "?")
                    updated.append({"id": str(p.id), "source": src})
            if next_offset is None:
                break
            offset = next_offset
        return {"success": True, "updated_count": len(updated), "updated": updated}
    except Exception as e:
        logger.error(f"Qdrant cleanup error: {e}")
        return JSONResponse({"success": False, "error": str(e)}, status_code=500)


# =========================
# RUN SERVER
# =========================

if __name__ == "__main__":

    import uvicorn

    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8096,
        reload=True
    )